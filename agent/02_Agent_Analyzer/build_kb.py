"""
RAG 知识库构建脚本 —— 支持 txt / md / docx / xlsx 格式，向量化存入 ChromaDB。
"""

import os
import sys
import re
import hashlib
from pathlib import Path
from typing import Optional

# Ensure parent `agent` directory is on sys.path for shared config
sys.path.append(str(Path(__file__).resolve().parents[1]))
from shared import PROJECT_ROOT, SOP_DOCS_DIR

import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer

# ─── 可配置项 ───────────────────────────────────────────────
EMBEDDING_MODEL = "BAAI/bge-small-zh-v1.5"
CHUNK_SIZE = 500
CHUNK_OVERLAP = 80
COLLECTION_NAME = "sop_methods"

# 知识库源目录（支持从环境变量 EXTERNAL_KNOWLEDGE_SOURCES 获取，以逗号分隔，并 fallback 默认值）
SOURCE_DIRS = []
env_sources = os.environ.get("EXTERNAL_KNOWLEDGE_SOURCES", "")
if env_sources:
    for path_str in env_sources.split(","):
        SOURCE_DIRS.append(Path(path_str.strip()))
else:
    SOURCE_DIRS = [Path("D:/香港房产账号资料")]
# ────────────────────────────────────────────────────────────

SCRIPT_DIR = Path(__file__).resolve().parent
LOCAL_SOP_DIR = SOP_DOCS_DIR
KB_DIR = SCRIPT_DIR / "knowledge_base"


# ══════════════════════════════════════════════════════════════
#  文档加载器
# ══════════════════════════════════════════════════════════════

def _load_txt(filepath: Path) -> Optional[str]:
    try:
        text = filepath.read_text(encoding="utf-8").strip()
        return text if text else None
    except UnicodeDecodeError:
        text = filepath.read_text(encoding="gbk").strip()
        return text if text else None


def _load_docx(filepath: Path) -> Optional[str]:
    try:
        from docx import Document
    except ImportError:
        print("[warn] 缺少 python-docx 库，跳过 .docx 文件。请执行: pip install python-docx")
        return None
    doc = Document(str(filepath))
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts) if parts else None


def _load_xlsx(filepath: Path) -> Optional[str]:
    try:
        import openpyxl
    except ImportError:
        print("[warn] 缺少 openpyxl 库，跳过 .xlsx 文件。请执行: pip install openpyxl")
        return None
    wb = openpyxl.load_workbook(str(filepath), data_only=True)
    all_parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        all_parts.append(f"# Sheet: {sheet_name}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                all_parts.append(" | ".join(cells))
    return "\n".join(all_parts) if all_parts else None


# ══════════════════════════════════════════════════════════════
#  文档分块
# ══════════════════════════════════════════════════════════════

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """按空行 → 句子边界递归分块。"""
    paragraphs = re.split(r"\n{2,}", text)
    chunks = []
    current = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) <= chunk_size:
            current = (current + "\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            if len(para) > chunk_size:
                sentences = re.split(r"(?<=[。！？；\n])", para)
                current = ""
                for sent in sentences:
                    sent = sent.strip()
                    if not sent:
                        continue
                    if len(current) + len(sent) <= chunk_size:
                        current = (current + sent).strip()
                    else:
                        if current:
                            chunks.append(current)
                        current = current[-overlap:] if current and len(current) > overlap else ""
                        current = current + sent
            else:
                current = para
    if current:
        chunks.append(current)
    return chunks


# ══════════════════════════════════════════════════════════════
#  主流程
# ══════════════════════════════════════════════════════════════

EXTENSION_LOADERS = {
    ".txt": _load_txt,
    ".md": _load_txt,
    ".docx": _load_docx,
    ".xlsx": _load_xlsx,
}


def collect_all_source_dirs() -> list[Path]:
    """汇总所有需要扫描的源目录。"""
    dirs = []
    for d in SOURCE_DIRS:
        if d.exists():
            dirs.append(d)
        else:
            print(f"[warn] 源目录不存在，跳过: {d}")
    if LOCAL_SOP_DIR.exists():
        dirs.append(LOCAL_SOP_DIR)
    return dirs


def load_all_documents(source_dirs: list[Path]) -> list[dict]:
    """遍历所有源目录，用对应的 loader 加载文件内容。"""
    docs = []
    for src_dir in source_dirs:
        print(f"[scan] 扫描目录: {src_dir}")
        for filepath in sorted(src_dir.iterdir()):
            if filepath.is_dir():
                continue
            if filepath.name.startswith("~$") or filepath.name.startswith("."):
                print(f"   [skip] temp/hidden: {filepath.name}")
                continue
            ext = filepath.suffix.lower()
            loader = EXTENSION_LOADERS.get(ext)
            if loader is None:
                print(f"   [skip] 不支持格式: {filepath.name}")
                continue
            print(f"   [load] {filepath.name} ...")
            content = loader(filepath)
            if content:
                docs.append({"path": str(filepath), "content": content, "source_name": filepath.stem})
                print(f"          -> {len(content)} 字")
            else:
                print(f"          -> 无内容/加载失败")
    print(f"\n[load] 共加载 {len(docs)} 个文档")
    return docs


def build_knowledge_base():
    print("=" * 60)
    print("  Agent_Analyzer — RAG 知识库构建")
    print("=" * 60)

    source_dirs = collect_all_source_dirs()
    if not source_dirs:
        print("[error] 没有可用的源目录，请在 SOURCE_DIRS 中配置。")
        sys.exit(1)
    print(f"[info] 源目录: {[str(d) for d in source_dirs]}\n")

    docs = load_all_documents(source_dirs)
    if not docs:
        print("[error] 未加载到任何文档内容。")
        sys.exit(1)

    all_chunks = []
    all_metadatas = []
    all_ids = []
    for doc in docs:
        chunks = chunk_text(doc["content"])
        for i, chunk in enumerate(chunks):
            chunk_id = hashlib.md5(f"{doc['source_name']}_{i}".encode()).hexdigest()[:12]
            all_chunks.append(chunk)
            all_metadatas.append({
                "source": doc["source_name"],
                "source_path": doc["path"],
                "chunk_index": i,
            })
            all_ids.append(chunk_id)
    print(f"[chunk] 共生成 {len(all_chunks)} 个文本块 (chunk_size={CHUNK_SIZE}, overlap={CHUNK_OVERLAP})")

    print(f"[embed] 加载模型: {EMBEDDING_MODEL}")
    model = SentenceTransformer(EMBEDDING_MODEL)
    print(f"[embed] 正在向量化 {len(all_chunks)} 个文本块 ...")
    embeddings = model.encode(all_chunks, normalize_embeddings=True, show_progress_bar=True)
    print(f"[embed] 向量维度: {embeddings.shape[1]}")

    KB_DIR.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(
        path=str(KB_DIR),
        settings=Settings(anonymized_telemetry=False),
    )
    try:
        client.delete_collection(COLLECTION_NAME)
        print(f"[db] 已删除旧 collection: {COLLECTION_NAME}")
    except Exception:
        pass

    collection = client.create_collection(
        name=COLLECTION_NAME,
        metadata={"description": "香港房产账号运营知识库", "hnsw:space": "cosine"},
    )
    collection.add(
        ids=all_ids,
        documents=all_chunks,
        embeddings=[emb.tolist() for emb in embeddings],
        metadatas=all_metadatas,
    )

    print(f"[db] 已写入 {collection.count()} 条记录到 collection '{COLLECTION_NAME}'")
    print(f"[db] 持久化路径: {KB_DIR}")
    print("=" * 60)
    print("  知识库构建完成！")
    print("=" * 60)


if __name__ == "__main__":
    build_knowledge_base()
